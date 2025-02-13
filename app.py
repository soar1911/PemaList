from flask import Flask, request, jsonify, render_template, send_file
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.shared import Length
import os
from datetime import datetime
from openpyxl import Workbook
from openpyxl.styles import Border, Side, Alignment, Font

app = Flask(__name__)

# 將常量移到配置類中
class Config:
    # 修改檔案輸出目錄使用 os.path.join
    OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'output_files')

    # 文件排版相關設定
    LINES_PER_PAGE = 30  # 每頁行數
    START_LINE = 20      # 起始行數
    MAX_CHARS_PER_LINE = 200  # 每行最大字元數

    # 字型大小設定（單位：pt）
    FONT_SIZES = {
        'title': 24,    # 標題字型大小
        'header': 14,   # 表頭字型大小
        'content': 14,  # 內容字型大小
        'normal': 12    # 一般文字字型大小
    }

    # 修改字型設定，增加 Linux 相容的字型
    FONT_NAMES = {
        'windows': 'Microsoft JhengHei',
        'linux': 'Noto Sans CJK TC'  # Linux 系統常用的中文字型
    }

    # 直接設置 FONT_NAME 為靜態屬性
    FONT_NAME = FONT_NAMES['windows'] if os.name == 'nt' else FONT_NAMES['linux']

    # Excel 欄寬設定
    COLUMN_WIDTHS = {
        'number': 10,  # 項次欄寬
        'name': 15,    # 姓名欄寬
        'book': 15,    # 法本欄寬
        'meal': 15,    # 便當欄寬
        'note': 30     # 備註欄寬
    }

    # Excel 列高設定
    ROW_HEIGHTS = {
        'title': 40,    # 標題列高
        'content': 20   # 內容列高
    }

    # 文件命名模板
    FILE_NAMES = {
        'xiazai': '消災牌位',
        'chaojian': '超薦牌位',
        'gongde': '功德主',
        'participant': '全程參加者名單'
    }

# 確保輸出目錄存在
if not os.path.exists(Config.OUTPUT_DIR):
    os.makedirs(Config.OUTPUT_DIR)

def ensure_output_directory():
    """確保輸出目錄存在且具有正確的權限"""
    if not os.path.exists(Config.OUTPUT_DIR):
        try:
            os.makedirs(Config.OUTPUT_DIR, mode=0o755)  # 設置目錄權限為 755
        except Exception as e:
            app.logger.error(f"無法創建輸出目錄: {str(e)}")
            raise

# 新增函數：找到符合條件的欄位名稱
def find_matching_column(df, keywords):
    """
    在DataFrame的欄位中尋找包含指定關鍵字的欄位名稱
    keywords 可以是單個字串或字串列表
    返回找到的第一個匹配欄位名稱，如果沒找到返回None
    """
    if isinstance(keywords, str):
        keywords = [keywords]

    keywords = [k.lower() for k in keywords]

    for col in df.columns:
        col_lower = str(col).lower()
        if any(keyword in col_lower for keyword in keywords):
            return col
    return None


def set_document_orientation_and_font(doc, is_landscape=True):
    section = doc.sections[0]
    if is_landscape:
        # 設定為橫式
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    else:
        # 設定為直式
        section.orientation = WD_ORIENT.PORTRAIT

    # 設定邊界為 1.27cm
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)

    # 設定預設字型大小
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(Config.FONT_SIZES['normal'])
    font.name = Config.FONT_NAME

def set_paragraph_format(paragraph):
    """設置段落格式：無間距，固定行高16pt"""
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(0)  # 段落後間距為0
    paragraph_format.space_before = Pt(0)  # 段落前間距為0
    paragraph_format.line_spacing = Pt(16)  # 固定行高16pt
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # 設置為固定行高

def add_empty_lines(doc, count):
    for _ in range(count):
        paragraph = doc.add_paragraph('')
        set_paragraph_format(paragraph)

def estimate_line_count(text, max_chars_per_line):
    """估算文字會佔用幾行"""
    if not text:
        return 1
    return -(-len(text) // max_chars_per_line)  # 向上取整除法

def create_word_document(is_landscape=False):
    """創建並初始化 Word 文件"""
    doc = Document()
    set_document_orientation_and_font(doc, is_landscape)
    return doc

def create_content_file(df, column_name, file_type, timestamp, is_landscape=False, prefix=""):
    """通用函數用於創建內容文件"""
    if not column_name or not df[column_name].notna().any():
        return None

    doc = create_word_document(is_landscape)

    # 預先過濾有效數據，減少迭代次數
    valid_rows = df[df[column_name].notna()].copy()

    def format_row_content(row):
        """格式化每一行的內容"""
        name = row['姓名']
        # 使用 replace 替換換行符，避免在 f-string 中使用反斜線
        content = str(row[column_name]).replace('\n', ' ').replace('\r', '')
        separator = ' | ' if prefix else '\t'
        return f"{prefix}{name}{separator}{content}"

    # 使用 apply 函數和新的格式化函數
    valid_rows['content'] = valid_rows.apply(format_row_content, axis=1)

    # 一次性添加空行
    add_empty_lines(doc, Config.START_LINE - 1)
    current_line = Config.START_LINE

    # 處理每一行內容
    for content in valid_rows['content']:
        estimated_lines = estimate_line_count(content, Config.MAX_CHARS_PER_LINE)

        if current_line + estimated_lines > Config.LINES_PER_PAGE:
            doc.add_page_break()
            add_empty_lines(doc, Config.START_LINE - 1)
            current_line = Config.START_LINE

        paragraph = doc.add_paragraph(content)
        set_paragraph_format(paragraph)
        current_line += estimated_lines

    # 保存文件
    file_path = os.path.join(Config.OUTPUT_DIR, f'{file_type}_{timestamp}.docx')
    try:
        doc.save(file_path)
        # 設置檔案權限（在 Linux 環境中）
        if os.name != 'nt':
            os.chmod(file_path, 0o644)
    except Exception as e:
        app.logger.error(f"保存文件失敗: {str(e)}")
        raise

    return file_path

def create_gongde_file(df, column_mapping, timestamp):
    """創建功德主文件"""
    if not column_mapping['gongde'] or not df[column_mapping['gongde']].notna().any():
        return None

    doc = create_word_document(is_landscape=True)

    # 預先過濾有效數據，並確保包含管理者註記事項欄位
    columns_to_keep = ['姓名', 'Email', '行動電話', column_mapping['gongde']]
    if column_mapping.get('note'):  # 如果有管理者註記事項欄位，加入到要保留的欄位中
        columns_to_keep.append(column_mapping['note'])

    # 過濾有效數據
    valid_rows = df[df[column_mapping['gongde']].notna()][columns_to_keep].copy()

    # 設置文件標題
    heading = doc.add_heading('功德主', 0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 設置標題字型
    for run in heading.runs:
        run.font.name = Config.FONT_NAME
        run.font.size = Pt(Config.FONT_SIZES['title'])

    # 創建並設置表格
    table = doc.add_table(rows=len(valid_rows) + 1, cols=5)
    table.style = 'Table Grid'
    table.autofit = True

    # 表頭設置
    headers = ['姓名', 'Email', '行動電話', '功德主', '管理者註記事項']
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = header
        # 設置表頭字型
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in paragraph.runs:
                run.font.name = Config.FONT_NAME
                run.font.size = Pt(Config.FONT_SIZES['header'])

    # 批量填充數據
    for row_idx, row in enumerate(valid_rows.iterrows(), 1):
        row_data = row[1]  # 獲取行數據
        cells = table.rows[row_idx].cells

        # 填充基本資料
        cells[0].text = str(row_data['姓名'])
        cells[1].text = str(row_data['Email'])
        cells[2].text = str(row_data['行動電話'])
        cells[3].text = str(row_data[column_mapping['gongde']])

        # 處理管理者註記事項
        if column_mapping.get('note'):
            note_value = row_data[column_mapping['note']]
            cells[4].text = str(note_value) if pd.notna(note_value) else ''
        else:
            cells[4].text = ''

        # 設置單元格對齊和字型
        for cell in cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.name = Config.FONT_NAME
                    run.font.size = Pt(Config.FONT_SIZES['content'])

    file_path = os.path.join(Config.OUTPUT_DIR, f'{Config.FILE_NAMES["gongde"]}_{timestamp}.docx')
    doc.save(file_path)
    return file_path

def create_word_files(df, column_mapping):  # 修改函數參數，接收 column_mapping
    """主函數：創建所有 Word 文件"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    file_paths = {}

    # 移除重複的 column_mapping 處理
    # 檢查是否成功匹配到必要的欄位
    if not column_mapping['gongde']:
        app.logger.warning('功德主欄位未找到，跳過功德主文件生成')

    # 創建消災牌位文件
    if path := create_content_file(df, column_mapping['xiazai'], Config.FILE_NAMES['xiazai'], timestamp):
        file_paths[Config.FILE_NAMES['xiazai']] = path

    # 創建超薦牌位文件
    if path := create_content_file(df, column_mapping['chaojian'], Config.FILE_NAMES['chaojian'], timestamp, prefix="陽上："):
        file_paths[Config.FILE_NAMES['chaojian']] = path

    # 創建功德主文件
    if column_mapping['gongde']:  # 確保功德主欄位存在
        if path := create_gongde_file(df, column_mapping, timestamp):
            file_paths[Config.FILE_NAMES['gongde']] = path

    return (
        file_paths,
        bool(column_mapping['xiazai']),
        bool(column_mapping['chaojian']),
        bool(column_mapping['gongde'])
    )

@app.route('/')
def index():
    return render_template('index.html')

def create_participant_excel(df):
    """創建參加者 Excel 檔案"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    # 預定義樣式（移到最前面）
    styles = {
        'title': Font(size=Config.FONT_SIZES['title'], name=Config.FONT_NAME),
        'header': Font(size=Config.FONT_SIZES['header'], name=Config.FONT_NAME),
        'content': Font(size=Config.FONT_SIZES['content'], name=Config.FONT_NAME),
        'border': Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
    }

    # 找到相關欄位並驗證
    columns = {
        'number': find_matching_column(df, '項次'),
        'name': find_matching_column(df, '姓名'),
        'activity': find_matching_column(df, ['參加項目', '參與課程']),
        'note': find_matching_column(df, '管理者註記事項')
    }

    if not all([columns['name'], columns['activity']]):
        return None

    # 使用向量化操作篩選參加者
    attendance_mask = df[columns['activity']].str.contains('現場上課|到場參加', case=False, na=False)
    participants = df[attendance_mask].copy()

    if participants.empty:
        return None

    # 創建工作簿和設置基本屬性
    wb = Workbook()
    ws = wb.active
    ws.title = "全程參加者名單"

    # 添加標題並設置格式
    title_cell = ws.cell(row=1, column=1, value="標題：現場名單")
    title_cell.font = styles['title']
    ws.merge_cells('A1:E1')
    title_cell.alignment = Alignment(horizontal='left', vertical='center')
    title_cell.border = styles['border']  # 添加邊框

    # 設置欄寬
    ws.column_dimensions['A'].width = Config.COLUMN_WIDTHS['number']
    ws.column_dimensions['B'].width = Config.COLUMN_WIDTHS['name']
    ws.column_dimensions['C'].width = Config.COLUMN_WIDTHS['book']
    ws.column_dimensions['D'].width = Config.COLUMN_WIDTHS['meal']
    ws.column_dimensions['E'].width = Config.COLUMN_WIDTHS['note']

    # 設置列高
    ws.row_dimensions[1].height = Config.ROW_HEIGHTS['title']
    for row in range(2, len(participants) + 3):
        ws.row_dimensions[row].height = Config.ROW_HEIGHTS['content']

    # 寫入表頭
    headers = ['項次', '姓名', '法本', '便當', '管理者註記事項']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=2, column=col, value=header)
        cell.font = styles['header']
        cell.border = styles['border']
        cell.alignment = Alignment(horizontal='center', vertical='center')

    # 準備並寫入數據
    for row_idx, (_, row) in enumerate(participants.iterrows(), 3):
        # 項次
        ws.cell(row=row_idx, column=1, value=row_idx-2).font = styles['content']
        # 姓名
        ws.cell(row=row_idx, column=2, value=row[columns['name']]).font = styles['content']
        # 法本（空白）
        ws.cell(row=row_idx, column=3, value='').font = styles['content']
        # 便當（空白）
        ws.cell(row=row_idx, column=4, value='').font = styles['content']
        # 備註
        note_value = row[columns['note']] if columns['note'] and pd.notna(row[columns['note']]) else ''
        ws.cell(row=row_idx, column=5, value=note_value).font = styles['content']

        # 設置每個單元格的邊框和對齊方式
        for col in range(1, 6):
            cell = ws.cell(row=row_idx, column=col)
            cell.border = styles['border']
            cell.alignment = Alignment(horizontal='center', vertical='center')

    # 儲存文件
    file_path = os.path.join(Config.OUTPUT_DIR, f'{Config.FILE_NAMES["participant"]}_{timestamp}.xlsx')
    wb.save(file_path)
    return file_path

@app.route('/process_excel', methods=['POST'])
def process_excel():
    """處理上傳的 Excel 檔案並生成相應文件"""

    def validate_file():
        """驗證上傳的文件"""
        if 'file' not in request.files:
            raise ValueError('沒有檔案')
        file = request.files['file']
        if file.filename == '':
            raise ValueError('沒有選擇檔案')
        return file

    def read_excel_file(file):
        """讀取並預處理 Excel 檔案"""
        df = pd.read_excel(file, dtype={'行動電話': str})
        if '行動電話' in df.columns:
            # 使用 vectorized 操作替代 apply
            mask = df['行動電話'].notna()
            df.loc[mask, '行動電話'] = df.loc[mask, '行動電話'].str.zfill(10)
        return df

    def validate_columns(df):
        """驗證必要欄位"""
        required_base_columns = {'姓名', 'Email', '行動電話'}  # 使用 set 提高查找效率
        missing_columns = required_base_columns - set(df.columns)
        if missing_columns:
            raise ValueError(f'缺少基本欄位：{", ".join(missing_columns)}')

    def validate_activity_data(df, activity_type):
        """驗證活動相關數據"""
        if activity_type != 'both':
            return None, None, None

        # 使用字典存儲列名和數據狀態，減少重複計算
        columns = {
            'xiazai': find_matching_column(df, '祈福牌位'),
            'chaojian': find_matching_column(df, ['超薦牌位', '超渡牌位']),
            'gongde': find_matching_column(df, '功德主')
        }

        if not any(columns.values()):
            raise ValueError('必須至少包含「祈福牌位」、「超薦牌位（或超渡牌位）」或「功德主」其中一個相關欄位')

        # 一次性檢查所有欄位的數據
        has_data = {
            key: col is not None and df[col].notna().any()
            for key, col in columns.items()
        }

        if not any(has_data.values()):
            raise ValueError('必須至少填寫「祈福牌位」、「超薦牌位」或「功德主」其中一項資料')

        return columns, has_data

    try:
        app.logger.info('開始處理上傳檔案')

        # 確保輸出目錄存在
        os.makedirs(Config.OUTPUT_DIR, exist_ok=True)

        response_data = {'message': '處理完成', 'files': {}}
        activity_type = request.form.get('activityType', 'both')

        # 驗證並讀取文件
        file = validate_file()
        df = read_excel_file(file)
        validate_columns(df)

        # 驗證活動數據
        columns, has_data = validate_activity_data(df, activity_type)

        # 處理法會相關文件
        if activity_type == 'both':
            try:
                # 找到管理者註記事項欄位
                note_column = find_matching_column(df, '管理者註記事項')

                # 創建一個包含所有必要數據的 DataFrame 視圖，避免重複訪問
                relevant_columns = ['姓名', 'Email', '行動電話']

                # 添加功能相關欄位
                for col in columns.values():
                    if col:
                        relevant_columns.append(col)

                # 添加管理者註記事項欄位
                if note_column:
                    relevant_columns.append(note_column)

                # 確保欄位名稱不重複
                relevant_columns = list(dict.fromkeys(relevant_columns))

                # 建立工作用的 DataFrame
                working_df = df[relevant_columns].copy()

                # 建立 column_mapping（只在這裡處理一次）
                column_mapping = {
                    'xiazai': find_matching_column(working_df, '祈福牌位'),
                    'chaojian': find_matching_column(working_df, ['超薦牌位', '超渡牌位']),
                    'gongde': find_matching_column(working_df, '功德主'),
                    'note': note_column
                }

                # 將 column_mapping 傳遞給 create_word_files
                file_paths, has_xiazai, has_chaojian, has_gongde = create_word_files(working_df, column_mapping)

                # 使用字典推導式簡化文件路徑處理
                file_types = {
                    'xiazai': '消災牌位',
                    'chaojian': '超薦牌位',
                    'gongde': '功德主'
                }

                for key, file_type in file_types.items():
                    if has_data.get(key) and file_type in file_paths:
                        response_data['files'][key] = os.path.basename(file_paths[file_type])

            except Exception as e:
                app.logger.error(f'Word 檔案創建失敗: {str(e)}')
                raise

        # 創建參加者名單
        try:
            if participant_excel := create_participant_excel(df):
                response_data['files']['participant'] = os.path.basename(participant_excel)
        except Exception as e:
            app.logger.error(f'參加者名單 Excel 創建失敗: {str(e)}')
            # 不中斷處理，繼續執行

        app.logger.info('處理完成')
        return jsonify(response_data), 200

    except ValueError as e:
        app.logger.error(str(e))
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        app.logger.error(f'處理過程發生錯誤：{str(e)}')
        return jsonify({'error': str(e)}), 500


@app.route('/download/<filename>')
def download_file(filename):
    try:
        if not os.path.exists(Config.OUTPUT_DIR):
            return jsonify({'error': '輸出目錄不存在'}), 400

        file_path = os.path.join(Config.OUTPUT_DIR, filename)
        if not os.path.exists(file_path):
            return jsonify({'error': '檔案不存在'}), 404

        return send_file(
            file_path,
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        return str(e), 400

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=33080, debug=True)
